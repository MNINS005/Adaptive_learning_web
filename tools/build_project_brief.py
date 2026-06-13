from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "adaptive_learning_project_brief.docx"
PDF_PATH = OUT_DIR / "adaptive_learning_project_brief.pdf"


SECTIONS = [
    (
        "Project Snapshot",
        [
            "Adaptive Learning Web is a plain Django learning platform for DSA practice. It tracks users, questions, attempts, and topic-wise skill state, then recommends a next question based on learner progress.",
            "The current deploy-friendly version runs without heavy ML dependencies by default. The ML recommender is optional and can be enabled with ENABLE_ML_RECOMMENDER=1 when model artifacts and ML packages are available.",
        ],
    ),
    (
        "Core Architecture",
        [
            "core: Django project settings, URL routing, WSGI/ASGI entry points.",
            "users: learner registration, login, profile views, and User ORM model.",
            "questions: question bank, next-question recommendation view, LeetCode CSV import command, and Question ORM model.",
            "attempts: attempt logging form, skill update logic, and Attempt ORM model.",
            "learning: dashboard, progress, knowledge state, retraining entry point, and KnowledgeState ORM model.",
            "templates/static: server-rendered HTML with template inheritance and a shared CSS file.",
        ],
    ),
    (
        "Database and ORM",
        [
            "Django ORM models map to PostgreSQL tables: users, questions, attempts, and knowledge_states.",
            "Foreign keys connect attempts to users and questions; knowledge states connect a user to a topic and skill score.",
            "Migrations create and update schema. You should not manually create tables with SQL for normal development.",
            "Local development can use local PostgreSQL through .env. Render production uses DATABASE_URL.",
        ],
    ),
    (
        "DSA Question Loading",
        [
            "The dataset notebook/leetcode_dataset - lc.csv contains title, description, difficulty, related_topics, and url columns.",
            "A management command now imports or updates questions through the Django ORM: python manage.py import_leetcode_questions --limit 300",
            "After import, verify URL coverage with: python manage.py shell -c \"from questions.models import Question; print(Question.objects.count(), Question.objects.exclude(leetcode_url__isnull=True).exclude(leetcode_url='').count())\"",
            "The Django admin question list now displays the leetcode_url field so you can visually confirm links are present.",
        ],
    ),
    (
        "Recommendation Flow",
        [
            "The fallback recommender selects a question from the learner's weakest topic when available, otherwise the first question.",
            "The optional ML recommender loads artifacts/dkt_model/dkt_model.keras and artifacts/rl_policy/rl_policy.pkl.",
            "DKT estimates a user's knowledge state from attempt history; the RL policy scores candidate questions and picks the best next problem.",
            "Render defaults to fallback mode to avoid TensorFlow build issues. Enable ML only after deploying model artifacts and ML dependencies.",
        ],
    ),
    (
        "Deployment on Render",
        [
            "Use build command: ./build.sh",
            "Use start command: gunicorn core.wsgi:application",
            "Set PYTHON_VERSION=3.12.5, SECRET_KEY, and DATABASE_URL in Render environment variables.",
            "Use the external PostgreSQL URL if the database is in a different Render account, and include sslmode=require.",
            "Static files are served through WhiteNoise; collectstatic runs during build.",
        ],
    ),
    (
        "SDE Interview Talking Points",
        [
            "Explain why the project uses standard Django apps instead of a monolithic API file: separation of concerns and maintainability.",
            "Describe the ORM relationships and how migrations keep schema changes reproducible.",
            "Discuss the fallback recommender as graceful degradation: the product still works when ML artifacts are unavailable.",
            "Mention production readiness improvements: environment variables, Gunicorn, WhiteNoise, Render PostgreSQL, and no secrets in code.",
            "Call out tradeoffs: server-rendered Django templates are simpler for this project, while a future React frontend could consume JSON views.",
        ],
    ),
    (
        "DS/ML Interview Talking Points",
        [
            "The DKT model represents learner knowledge as a sequence model over past attempts.",
            "Input sequences encode question identity and correctness, allowing the model to infer topic mastery trends.",
            "The RL policy treats question selection as a decision problem and rewards appropriate challenge levels.",
            "The system combines interpretable rule-based updates to KnowledgeState with optional model-based recommendation.",
            "Production ML concern: TensorFlow is heavy on small hosts, so model loading is made optional and isolated behind an environment flag.",
        ],
    ),
]


CHECKLIST = [
    ["Area", "What to verify"],
    ["Question data", "Run import_leetcode_questions and confirm leetcode_url count is greater than zero."],
    ["Local DB", "PostgreSQL service is running and .env has DB_NAME, DB_USER, DB_PASSWORD, DB_HOST, DB_PORT."],
    ["Render DB", "DATABASE_URL uses the Render external URL with sslmode=require."],
    ["ML mode", "ENABLE_ML_RECOMMENDER is unset/0 unless TensorFlow, Keras, NumPy, artifacts, and URLs are deployed."],
]


def add_docx_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181 if level == 1 else 120)
    return para


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Adaptive Learning Web")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Interview Brief for SDE and DS/ML Discussions").italic = True

    for heading, bullets in SECTIONS:
        add_docx_heading(doc, heading, 1)
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    add_docx_heading(doc, "Pre-Demo Checklist", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = CHECKLIST[0][0]
    header[1].text = CHECKLIST[0][1]
    for area, detail in CHECKLIST[1:]:
        row = table.add_row().cells
        row[0].text = area
        row[1].text = detail

    doc.save(DOCX_PATH)


def pdf_header(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 9)
    canvas.setFillColor(colors.HexColor("#667085"))
    canvas.drawString(inch, 0.5 * inch, "Adaptive Learning Web - Interview Brief")
    canvas.drawRightString(7.5 * inch, 0.5 * inch, f"Page {doc.page}")
    canvas.restoreState()


def bullet_list(items, styles):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["Body"])) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=18,
    )


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="TitleBlue",
            parent=styles["Title"],
            textColor=colors.HexColor("#1F4D78"),
            fontSize=22,
            leading=26,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Subtitle",
            parent=styles["BodyText"],
            alignment=1,
            textColor=colors.HexColor("#475467"),
            fontSize=11,
            leading=14,
            spaceAfter=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1Blue",
            parent=styles["Heading1"],
            textColor=colors.HexColor("#2E74B5"),
            fontSize=15,
            leading=18,
            spaceBefore=12,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontSize=10.5,
            leading=13,
            spaceAfter=4,
        )
    )

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.9 * inch,
        bottomMargin=0.8 * inch,
    )
    story = [
        Paragraph("Adaptive Learning Web", styles["TitleBlue"]),
        Paragraph("Interview Brief for SDE and DS/ML Discussions", styles["Subtitle"]),
    ]

    for heading, bullets in SECTIONS:
        story.append(Paragraph(heading, styles["H1Blue"]))
        story.append(bullet_list(bullets, styles))
        story.append(Spacer(1, 6))

    story.append(PageBreak())
    story.append(Paragraph("Pre-Demo Checklist", styles["H1Blue"]))
    table = Table(CHECKLIST, colWidths=[1.55 * inch, 4.95 * inch], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C4D6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEADING", (0, 0), (-1, -1), 11),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 12))
    story.append(Paragraph("Command Reference", styles["H1Blue"]))
    story.append(
        bullet_list(
            [
                "Load questions: python manage.py import_leetcode_questions --limit 300",
                "Verify URLs: python manage.py shell -c \"from questions.models import Question; print(Question.objects.count(), Question.objects.exclude(leetcode_url__isnull=True).exclude(leetcode_url='').count())\"",
                "Run locally: python manage.py runserver",
                "Render deploy: build command ./build.sh, start command gunicorn core.wsgi:application",
            ],
            styles,
        )
    )

    doc.build(story, onFirstPage=pdf_header, onLaterPages=pdf_header)


def main():
    OUT_DIR.mkdir(exist_ok=True)
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
